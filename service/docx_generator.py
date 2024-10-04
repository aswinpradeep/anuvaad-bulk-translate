import json
import docx
from docx.shared import Pt
from docx.shared import RGBColor
import os
from docx.enum.text import WD_BREAK
from config.config import logger

# #delete existing output_doc.docx
# try:
#     os.remove("./output_doc.docx")
# except:
#     logger.debug("No file to be deleted")
    
#Read fetch_content

class Docx_Generator():

    """
    Things to Consider for later on
    Background Images
    Lines
    """ 

    """
    "s0_src": "Input"
    "s0_tgt": "TMX / UTM Output if it exists otherwise NMT Output"
    "src": "Input"
    "tgt": "TMX / UTM Output if it exists otherwise NMT Output"
    "tagged_src": "Input"
    "tagged_tgt": "NMT Output"
    """

    def hex_to_rgb(self,hex_color):
        # Remove the '#' symbol if present
        hex_color = hex_color.lstrip('#')
        # Convert the hexadecimal string to RGB values
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

    def add_text_with_position(self,doc, page_height, page_width, text_top, text_left, text, font_color, font_attrib, font_family, font_size, text_width):
        if text == "None" and font_size == None:
            return doc
        # Calculate the Y-coordinate based on page height and text top
        y = page_height - text_top
        
        # Create a paragraph for the text
        p = doc.add_paragraph()
        

        #Check if Bold / Italic attribute is present
        if font_attrib is not None:
            if "BOLD" in font_attrib:
                p.bold = True
            if "ITALIC" in font_attrib:
                p.italic = True

        # Calculate the number of spaces needed before the text
        # font_size = 11
        if font_size == None:
            font_size = 11
        # else:
        #     font_size = int(font_size*0.8)
        spacing_factor = 7.5 * (font_size / 12)  # Adjust 12 to your reference font size
        num_spaces = int(text_left / spacing_factor)
        spaces = " " * num_spaces

        p.style.font.size = Pt(12)            
        p.style.font.name = "Raavi"
        p.style.paragraph_format.line_spacing = 1
    
        # Add the spaces and text to the paragraph
        p.add_run(spaces + text)

        # Convert the hexadecimal font color to RGBColor
        if font_color is not None:
            font_color_rgb = RGBColor(*self.hex_to_rgb(font_color))
        else:
            font_color_rgb = RGBColor(*self.hex_to_rgb("#000000"))
        
        # Set font properties
        for run in p.runs:
            run.font.color.rgb = font_color_rgb  # RGB color

        # Set font properties

        # Calculate the text width and set a tab stop for the entire paragraph
        tab_stop_position = Pt(text_left + text_width)
        p.paragraph_format.tab_stops.add_tab_stop(tab_stop_position)
            
        # Set the vertical alignment of the paragraph
        p.alignment = 3  # 3 corresponds to centered alignment
        
        return doc

    def func(self, block):
        return block["text_top"]

    def generate_docx(self,fetch_content,output_filename):
        # fetch_content_file = "./fetch_content.json"
        # with open(fetch_content_file,"r") as file:
        # fetch_content = file.read()
        # fetch_content = json.loads(fetch_content)

        #Read common info
        page_height = fetch_content["data"][0]["page_height"]
        page_width = fetch_content["data"][0]["page_width"]

        #Create a document
        document = docx.Document()


        for block in fetch_content["data"]:
            if "text_blocks" in block.keys():
                text_blocks = block["text_blocks"].copy()
                text_blocks.sort(key=self.func)
                merged_text_blocks = []

                i=0
                while i<len(text_blocks):
                    current_block = text_blocks[i]
                    merged_block = current_block.copy()
                    # Check if there are subsequent blocks with the same 'text_left'
                    while i + 1 < len(text_blocks) and text_blocks[i]["text_left"] == text_blocks[i + 1]["text_left"]:
                        spacing_factor = 7.5 * (11 / 12)  # Adjust 12 to your reference font size
                        num_spaces = int(text_blocks[i+1]["text_left"]-text_blocks[i+1]["text_left"]/ spacing_factor)
                        current_text = [token["s0_tgt"] for token in text_blocks[i]["tokenized_sentences"]]
                        next_text = [token["s0_tgt"] for token in text_blocks[i+1]["tokenized_sentences"]]
                        current_text = " ".join(current_text)
                        next_text = " ".join(next_text)
                        merged_block["merged_text"] = current_text + " "*num_spaces + next_text  # Merge text
                        # You can also merge other properties if needed
                        i += 1

                    merged_text_blocks.append(merged_block)
                    i += 1


                for each_block in merged_text_blocks:
                    #Get block info
                    font_attrib = each_block["attrib"]
                    font_color = each_block["font_color"]
                    font_family = each_block["font_family"]
                    font_size = each_block["font_size"]
                    text_height = each_block["text_height"]
                    text_left = each_block["text_left"]
                    text_top = each_block["text_top"]
                    text_width = each_block["text_width"]
                    if "merged_text" in each_block.keys():
                        text = each_block["merged_text"]
                    else:
                        text = [token["s0_tgt"] for token in each_block["tokenized_sentences"]]
                        text = " ".join(text)
                    #logger.debug(text)
                    self.add_text_with_position(document,page_height,page_width,text_top,text_left, text,
                                        font_color,font_attrib, font_family, font_size, text_width)
                    #break
            p = document.add_paragraph()
            run = p.add_run()
            run.add_break(WD_BREAK.PAGE)
        logger.debug(f"Document Saved :: {output_filename}")
        document.save(output_filename)